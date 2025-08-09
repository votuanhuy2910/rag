import streamlit as st
import pdfplumber
import docx
import os
import google.generativeai as genai
from sentence_transformers import SentenceTransformer
import chromadb
import csv

# ===== CONFIG GEMINI KEY =====
# Điền key Gemini của bạn vào đây ↓↓↓
GEMINI_API_KEY = "AIzaSyCSenmJGRf2VJ9WId1SwQpfL3dMRRaHWmw"
genai.configure(api_key=GEMINI_API_KEY)

# Load embedding model
embedding_model = SentenceTransformer('sentence-transformers/all-MiniLM-L6-v2')

# Chroma vector DB setup
chroma_client = chromadb.Client()
collection = chroma_client.get_or_create_collection("essays")


# Hàm export collection ra CSV dự phòng
def export_collection_to_csv(collection, file_path="essays_backup.csv"):
    all_data = collection.get()
    with open(file_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["id", "document", "metadata"])
        for i in range(len(all_data["ids"])):
            metadata = all_data["metadatas"][i] if "metadatas" in all_data else {}
            writer.writerow([
                all_data["ids"][i],
                all_data["documents"][i],
                metadata
            ])

def load_sample_data():
    sample_folder = "data"
    if not os.path.exists(sample_folder):
        os.makedirs(sample_folder)

    docs, ids = [], []

    for idx, filename in enumerate(os.listdir(sample_folder)):
        path = os.path.join(sample_folder, filename)

        text = ""
        if filename.endswith(".txt"):
            with open(path, "r", encoding="utf-8") as f:
                text = f.read()

        elif filename.endswith(".pdf"):
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

        elif filename.endswith(".docx"):
            import docx
            doc = docx.Document(path)
            text = "\n".join([para.text for para in doc.paragraphs])

        else:
            continue  # bỏ qua file không hỗ trợ

        if text.strip():
            docs.append(text)
            ids.append(str(idx))

    # Thêm vào vector DB
    if docs:
        embeddings = embedding_model.encode(docs).tolist()
        collection.add(documents=docs, ids=ids, embeddings=embeddings)
        export_collection_to_csv(collection, "essays_backup.csv")  # Export sau khi thêm

# Load ngay khi khởi động app
load_sample_data()


# ===== FUNCTIONS =====
def read_file(file):
    if file.name.endswith(".pdf"):
        text = ""
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + "\n"
        return text
    elif file.name.endswith(".docx"):
        doc = docx.Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    elif file.name.endswith(".txt"):
        return file.read().decode("utf-8")
    else:
        return ""

def find_relevant_docs(query):
    query_embedding = embedding_model.encode([query]).tolist()
    results = collection.query(query_embeddings=query_embedding, n_results=1)
    return results['documents'][0][0] if results['documents'] else ""

def grade_essay(essay_text, course_context, sample_text):
    model = genai.GenerativeModel("gemini-1.5-flash")
    prompt = f"""
Bạn là giảng viên đại học. Hãy chấm bài luận của sinh viên theo thang điểm 10.

Hãy chấm điểm bài luận dưới đây dựa trên bộ tiêu chí sau: {rubric}
Yêu cầu:
1. Chấm từng tiêu chí kèm điểm số.
2. Tính tổng điểm cuối cùng.
3. Đưa ra nhận xét chi tiết, chỉ ra điểm mạnh và điểm cần cải thiện.

Ngữ cảnh môn học: {course_context}
Ngữ cảnh sinh viên: {student_context}
Bài mẫu tham chiếu: {sample_text}
Bài luận cần chấm: {essay_text}

Trả về:
- Điểm tổng (0-10)
- Nhận xét chi tiết
"""
    response = model.generate_content(prompt)
    return response.text

# ===== STREAMLIT UI =====
st.title("📄 Chấm điểm bài luận với Gemini + RAG")

uploaded_file = st.file_uploader("Tải lên bài luận (.pdf, .docx, .txt)", type=["pdf", "docx", "txt"])
course_context = st.selectbox("Chọn ngữ cảnh môn học", ["Thương mại điện tử", "Phát triển ứng dụng giao diện", "Hệ thống thanh toán trong thương mại điện tử", "Chiến lược thương mại điện tử"])
year = st.selectbox("Chọn năm học", ["Năm 1", "Năm 2", "Năm 3", "Năm 4"])
faculty = st.selectbox("Chọn ngành", ["Công nghệ thông tin", "Ngôn ngữ Anh", "Quản tri kinh doanh", "Kế toán", "Marketing"])
student_context = f"Sinh viên {year} thuộc Khoa {faculty}"
rubric = """
Tiêu chí chấm điểm thuộc lĩnh vực Công nghệ thông tin. Tổng điểm tối đa là 10 điểm, chia thành 5 nhóm tiêu chí chính:
1. Nội dung và kiến thức chuyên môn (4.5 điểm)
2. Cấu trúc và tổ chức bài viết (2 điểm)
3. Trình bày và diễn đạt (1.5 điểm)
4. Phân tích, lập luận và giải thích (1.5 điểm)
5. Tài liệu tham khảo và trích dẫn (0.5 điểm)
Ngoài ra có tiêu chí trừ điểm nếu bài làm vi phạm yêu cầu nghiêm trọng.

Tiêu chí 1 – Nội dung và kiến thức chuyên môn (4.5 điểm):
• Đáp ứng yêu cầu đề bài (1.5 điểm): bài làm đúng trọng tâm, không lạc đề, bao quát đủ các khía cạnh.
• Tính chính xác kiến thức (2 điểm): nội dung đúng, không sai nghiêm trọng về lý thuyết, thuật toán, phương pháp.
• Tính sáng tạo và thực tiễn (1 điểm): có tiếp cận mới, vận dụng thực tế, có dẫn chứng rõ.

Tiêu chí 2 – Cấu trúc và tổ chức bài viết (2 điểm):
• Cấu trúc (1.2 điểm): bài gồm mở bài, thân bài, kết luận; bố cục rõ ràng, không rời rạc.
• Logic liên kết (0.8 điểm): các ý nối liền mạch, có sử dụng từ nối để giữ tính liên tục.

Tiêu chí 3 – Trình bày và diễn đạt (1.5 điểm):
• Ngôn ngữ chuyên ngành (0.6 điểm): sử dụng đúng thuật ngữ CNTT, tránh cách nói đơn giản hóa.
• Chính tả, ngữ pháp (0.6 điểm): không mắc lỗi nghiêm trọng, diễn đạt mạch lạc.
• Hình thức trình bày (0.3 điểm): rõ ràng, khoa học, có số trang/mục lục nếu cần.

Tiêu chí 4 – Phân tích, lập luận và giải thích (1.5 điểm):
• Lập luận chặt chẽ (0.7 điểm): có cơ sở rõ ràng, giải thích hợp lý.
• Ví dụ, dẫn chứng (0.5 điểm): có minh họa bằng thực tế, sơ đồ hoặc thuật toán.
• Phản biện, đánh giá (0.3 điểm): đưa nhiều góc nhìn, so sánh giải pháp khác nhau.

Tiêu chí 5 – Tài liệu tham khảo và trích dẫn (0.5 điểm):
• Tài liệu chất lượng (0.3 điểm): có nguồn gốc rõ ràng, chính thống.
• Trích dẫn đúng chuẩn (0.2 điểm): theo quy chuẩn như APA, IEEE, Harvard...

Tiêu chí trừ điểm:
• Lạc đề hoặc sai nghiêm trọng về kiến thức: trừ toàn bộ điểm phần nội dung.
• Trình bày quá sơ sài, không có cấu trúc rõ ràng: trừ tối đa 2 điểm.
• Viết sai chính tả, ngữ pháp quá nhiều: trừ tối đa 1 điểm.
• Không trích nguồn nhưng dùng tài liệu ngoài: trừ 0.5 điểm.
"""


if st.button("Chấm điểm"):
    if uploaded_file is not None:
        essay_text = read_file(uploaded_file)
        if not essay_text.strip():
            st.error("Không đọc được nội dung từ file.")
        else:
            sample_text = find_relevant_docs(essay_text)
            result = grade_essay(essay_text, course_context, sample_text)
            st.subheader("Kết quả chấm điểm")
            st.write(result)
    else:
        st.warning("Vui lòng tải lên bài luận trước.")
