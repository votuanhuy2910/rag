import streamlit as st
import pdfplumber
import warnings
warnings.filterwarnings("ignore", category=UserWarning, module="pdfminer")
import docx
import re
import os
import google.generativeai as genai
from sentence_transformers import SentenceTransformer
import chromadb
import csv

# ===== CONFIG GEMINI KEY =====
# Điền key Gemini của bạn vào đây ↓↓↓
GEMINI_API_KEY = "AIzaSyCSenmJGRf2VJ9WId1SwQpfL3dMRRaHWmw"
genai.configure(api_key=GEMINI_API_KEY)

# ===== RUBRIC =====
rubric = """
Tiêu chí chấm điểm thuộc lĩnh vực Công nghệ thông tin. Tổng điểm tối đa là 10 điểm, chia thành 5 nhóm tiêu chí chính:
1. Nội dung và kiến thức chuyên môn (4.5 điểm)
2. Cấu trúc và tổ chức bài viết (2 điểm)
3. Trình bày và diễn đạt (1.5 điểm)
4. Phân tích, lập luận và giải thích (1.5 điểm)
5. Tài liệu tham khảo và trích dẫn (0.5 điểm)
Ngoài ra có tiêu chí trừ điểm nếu bài làm vi phạm yêu cầu nghiêm trọng.

Tiêu chí 1 – Nội dung và kiến thức chuyên môn (4.5 điểm):
• Đáp ứng yêu cầu đề bài (1.5 điểm): bài làm đúng trọng tâm, không lạc đề, bao quát đủ các khía cạnh.
• Tính chính xác kiến thức (2 điểm): nội dung đúng, không sai nghiêm trọng về lý thuyết, thuật toán, phương pháp.
• Tính sáng tạo và thực tiễn (1 điểm): có tiếp cận mới, vận dụng thực tế, có dẫn chứng rõ.

Tiêu chí 2 – Cấu trúc và tổ chức bài viết (2 điểm):
• Cấu trúc (1.2 điểm): bài gồm mở bài, thân bài, kết luận; bố cục rõ ràng, không rời rạc.
• Logic liên kết (0.8 điểm): các ý nối liền mạch, có sử dụng từ nối để giữ tính liên tục.

Tiêu chí 3 – Trình bày và diễn đạt (1.5 điểm):
• Ngôn ngữ chuyên ngành (0.6 điểm): sử dụng đúng thuật ngữ CNTT, tránh cách nói đơn giản hóa.
• Chính tả, ngữ pháp (0.6 điểm): không mắc lỗi nghiêm trọng, diễn đạt mạch lạc.
• Hình thức trình bày (0.3 điểm): rõ ràng, khoa học, có số trang/mục lục nếu cần.

Tiêu chí 4 – Phân tích, lập luận và giải thích (1.5 điểm):
• Lập luận chặt chẽ (0.7 điểm): có cơ sở rõ ràng, giải thích hợp lý.
• Ví dụ, dẫn chứng (0.5 điểm): có minh họa bằng thực tế, sơ đồ hoặc thuật toán.
• Phản biện, đánh giá (0.3 điểm): đưa nhiều góc nhìn, so sánh giải pháp khác nhau.

Tiêu chí 5 – Tài liệu tham khảo và trích dẫn (0.5 điểm):
• Tài liệu chất lượng (0.3 điểm): có nguồn gốc rõ ràng, chính thống.
• Trích dẫn đúng chuẩn (0.2 điểm): theo quy chuẩn như APA, IEEE, Harvard...

Tiêu chí trừ điểm:
• Lạc đề hoặc sai nghiêm trọng về kiến thức: trừ toàn bộ điểm phần nội dung.
• Trình bày quá sơ sài, không có cấu trúc rõ ràng: trừ tối đa 2 điểm.
• Viết sai chính tả, ngữ pháp quá nhiều: trừ tối đa 1 điểm.
• Không trích nguồn nhưng dùng tài liệu ngoài: trừ 0.5 điểm.
"""

# Load embedding model
embedding_model = SentenceTransformer('sentence-transformers/all-MiniLM-L6-v2')

# Chroma vector DB setup
chroma_client = chromadb.Client()
collection = chroma_client.get_or_create_collection("essays")

# dữ liệu nền => những file mẫu (pdf, docx, txt)

# Hàm export collection ra CSV dự phòng (UTF-8 with BOM, metadata tách cột)
def export_collection_to_csv(collection, file_path="essays_backup.csv"):
    all_data = collection.get()
    with open(file_path, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        # Header
        writer.writerow(["id", "document", "course", "category", "filename", "basename"])
        
        for i in range(len(all_data["ids"])):
            metadata = all_data["metadatas"][i] if "metadatas" in all_data else {}
            writer.writerow([
                all_data["ids"][i],
                all_data["documents"][i],
                metadata.get("course", ""),
                metadata.get("category", ""),
                metadata.get("filename", ""),
                metadata.get("basename", "")
            ])


def load_sample_data():
    base_folder = "data"

    docs, ids, metadatas = [], [], []

    # Duyệt qua từng môn học
    if not os.path.exists(base_folder):
        os.makedirs(base_folder)

    for course in os.listdir(base_folder):
        course_path = os.path.join(base_folder, course)
        if not os.path.isdir(course_path):
            continue  # bỏ qua file lẻ, chỉ lấy folder môn học

        # Các thư mục con: essays, scores, teaching_materials
        for category in ["essays", "scores", "teaching_materials"]:
            folder_path = os.path.join(course_path, category)
            if not os.path.exists(folder_path):
                os.makedirs(folder_path)  # tạo thư mục nếu chưa có

            for idx, filename in enumerate(os.listdir(folder_path)):
                path = os.path.join(folder_path, filename)
                text = ""

                # Đọc file txt
                if filename.endswith(".txt"):
                    with open(path, "r", encoding="utf-8") as f:
                        text = f.read()

                # Đọc file pdf
                elif filename.endswith(".pdf"):
                    with pdfplumber.open(path) as pdf:
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"

                # Đọc file docx
                elif filename.endswith(".docx"):
                    doc = docx.Document(path)
                    text = "\n".join([para.text for para in doc.paragraphs])

                else:
                    continue  # bỏ qua định dạng không hỗ trợ

                if text.strip():
                    basename = os.path.splitext(filename)[0]
                    docs.append(text)
                    ids.append(f"{course}_{category}_{idx}")
                    metadatas.append({
                        "course": course,
                        "category": category,
                        "filename": filename,
                        "basename": basename
                    })

    # Thêm vào vector DB
    if docs:
        embeddings = embedding_model.encode(docs).tolist()
        collection.add(documents=docs, ids=ids, embeddings=embeddings, metadatas=metadatas)
        export_collection_to_csv(collection, "essays_backup.csv")

# Load ngay khi khởi động app
load_sample_data()


# ===== FUNCTIONS =====
def read_file(file):
    if file.name.endswith(".pdf"):
        text = ""
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + "\n"
        return text
    elif file.name.endswith(".docx"):
        doc = docx.Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    elif file.name.endswith(".txt"):
        return file.read().decode("utf-8")
    else:
        return ""

def save_result_to_excel(course, filename, essay_text, score, file_path="grading_results.xlsx"):
    from openpyxl import Workbook, load_workbook
    
    #  Chỉ lấy 300 ký tự đầu tiên
    essay_preview = essay_text[:300] + ("..." if len(essay_text) > 300 else "")
    
    if os.path.exists(file_path):
        wb = load_workbook(file_path)
        ws = wb.active
    else:
        wb = Workbook()
        ws = wb.active
        ws.append(["Tên môn học", "Tên file", "Nội dung bài luận", "Điểm"])

    ws.append([course, filename, essay_preview, score])
    wb.save(file_path)


def find_relevant_docs(query):
    query_embedding = embedding_model.encode([query]).tolist()
    results = collection.query(query_embeddings=query_embedding, n_results=1)
    return results['documents'][0][0] if results['documents'] else ""

def grade_essay(essay_text, course_context, sample_text):
    model = genai.GenerativeModel("gemini-1.5-flash")
    prompt = f"""
Bạn là giảng viên đại học. Hãy chấm bài luận của sinh viên theo thang điểm 10.

Hãy chấm điểm bài luận dưới đây dựa trên bộ tiêu chí sau: {rubric}
Yêu cầu:
1. Chấm từng tiêu chí kèm điểm số.
2. Tính tổng điểm cuối cùng.
3. Đưa ra nhận xét chi tiết, chỉ ra điểm mạnh và điểm cần cải thiện.

Ngữ cảnh môn học: {course_context}
Ngữ cảnh sinh viên: {student_context}
Bài mẫu tham chiếu: {sample_text}
Bài luận cần chấm: {essay_text}

Trả về:
- Điểm tổng (0-10)
- Nhận xét chi tiết
"""
    response = model.generate_content(prompt)
    return response.text

def extract_score(result_text: str) -> str:
    import re
    score_value = "?"
    # Ưu tiên các dòng chứa từ khóa
    for line in result_text.splitlines():
        if any(kw in line for kw in ["Điểm tổng", "Tổng điểm", "Điểm cuối cùng"]):
            match = re.search(r"(\d+(\.\d+)?)", line)
            if match:
                score_value = match.group(1)
            break

    # Fallback: tìm số trong toàn bộ văn bản
    if score_value == "?":
        matches = re.findall(r"(\d+(\.\d+)?)", result_text)
        candidates = [float(m[0]) for m in matches]
        candidates = [c for c in candidates if 0 <= c <= 10]
        if candidates:
            score_value = str(max(candidates))

    return score_value


# ===== STREAMLIT UI =====
st.set_page_config(page_title="Chấm điểm bài luận", page_icon="📄", layout="wide")

# CSS
st.markdown("""
    <style>
    .main {
        background-color: #f9fafc;
    }
    .stTitle {
        font-size: 28px !important;
        color: #1a237e;
        font-weight: bold;
    }
    .score-box {
        padding: 20px;
        border-radius: 12px;
        background-color: #e3f2fd;
        text-align: center;
        margin-bottom: 20px;
    }
    .score-text {
        font-size: 48px;
        font-weight: bold;
        color: #0d47a1;
    }
    </style>
""", unsafe_allow_html=True)

st.title("📄 Hệ thống chấm điểm bài luận tự động")

# Layout
col1, col2 = st.columns([1, 2])

with col1:
    st.header("📂 Nhập dữ liệu")
    uploaded_file = st.file_uploader("Tải lên bài luận (.pdf, .docx, .txt)", type=["pdf", "docx", "txt"])
    course_context = st.selectbox("📘 Môn học", ["Thương mại điện tử", "Phát triển ứng dụng giao diện", "Hệ thống thanh toán trong thương mại điện tử", "Chiến lược thương mại điện tử"])
    year = st.selectbox("🎓 Năm học", ["Năm 1: Sinh viên chỉ mới làm quen viết luận. Đánh giá chủ yếu ở sự rõ ràng, logic cơ bản, cách trình bày ý tưởng. Không yêu cầu nhiều về trích dẫn học thuật hay cấu trúc phức tạp.", "Năm 2: Sinh viên bắt đầu học kỹ năng viết nâng cao hơn. Cần có cấu trúc 3 phần (mở bài – thân bài – kết luận), biết triển khai luận điểm theo đoạn văn mạch lạc, có ví dụ minh họa cơ bản.", "Năm 3: Sinh viên phải thể hiện lập luận chặt chẽ hơn, biết sử dụng tài liệu tham khảo (trích dẫn đúng cách), trình bày theo chuẩn học thuật, có phân tích, so sánh, đánh giá thay vì chỉ mô tả.", "Năm 4: Sinh viên cần đạt chuẩn luận văn tốt nghiệp: viết học thuật hoàn chỉnh, có đặt vấn đề – cơ sở lý thuyết – phương pháp – phân tích – kết quả – kết luận, sử dụng trích dẫn chuẩn, thể hiện tư duy nghiên cứu độc lập và đóng góp mới."])
    faculty = st.selectbox("🏫 Khoa", ["Công nghệ thông tin kinh doanh", "Quản trị kinh doanh", "Marketing"])
    student_context = f"Sinh viên {year} thuộc Khoa {faculty}"

    if st.button("🚀 Chấm điểm", use_container_width=True):
        if uploaded_file is not None:
            essay_text = read_file(uploaded_file)
            if not essay_text.strip():
                st.error("❌ Không đọc được nội dung từ file.")
            else:
                sample_text = find_relevant_docs(essay_text)
                result = grade_essay(essay_text, course_context, sample_text)
                st.session_state["grading_result"] = result
                
                score_value = extract_score(result)

                # Lưu kết quả vào Excel
                save_result_to_excel(course_context, uploaded_file.name, essay_text, score_value)
                st.success(f"✅ Kết quả đã được lưu vào grading_results.xlsx (Điểm: {score_value})")
        else:
            st.warning("⚠️ Vui lòng tải lên bài luận trước.")

with col2:
    st.header("📊 Kết quả chấm điểm")
    if "grading_result" in st.session_state:
        result_text = st.session_state["grading_result"]
        score_value = extract_score(result_text)
        st.markdown(f"""
            <div class="score-box">
                <div style="color: #000">Điểm tổng</div>
                <div class="score-text">{score_value}</div>
            </div>
        """, unsafe_allow_html=True)

        st.subheader("📝 Nhận xét chi tiết")
        st.write(result_text)
    else:
        st.info("👉 Kết quả sẽ hiển thị tại đây sau khi chấm điểm.")

# ===== QUẢN LÝ CORPUS =====
st.markdown("---")
st.header("📂 Quản lý Corpus")

if st.button("🔄 Làm mới dữ liệu corpus"):
    load_sample_data()
    st.success("Dữ liệu corpus đã được làm mới.")

all_data = collection.get()
if all_data and all_data["ids"]:
    import pandas as pd
    rows = []
    for i in range(len(all_data["ids"])):
        meta = all_data["metadatas"][i] if "metadatas" in all_data else {}
        rows.append({
            "id": all_data["ids"][i],
            "course": meta.get("course", ""),
            "category": meta.get("category", ""),
            "filename": meta.get("filename", ""),
            "basename": meta.get("basename", ""),
            "preview": all_data["documents"][i][:200] + "..."
        })
    df = pd.DataFrame(rows)
    st.dataframe(df, use_container_width=True)

    selected_id = st.selectbox("🔍 Xem nội dung chi tiết theo ID:", df["id"])
    if selected_id:
        idx = df.index[df["id"] == selected_id][0]
        st.subheader(f"📄 Nội dung chi tiết của `{df.at[idx, 'filename']}`")
        st.text_area("Document", all_data["documents"][idx], height=300)
else:
    st.info("📭 Chưa có dữ liệu corpus. Vui lòng thêm file vào thư mục `data/` và làm mới.")