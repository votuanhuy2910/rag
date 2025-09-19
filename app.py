import streamlit as st
import pdfplumber
import fitz
import warnings
warnings.filterwarnings("ignore", category=UserWarning, module="pdfminer")
import docx
import re
import os
import google.generativeai as genai
from sentence_transformers import SentenceTransformer
import chromadb
import csv

# =======================================================================
# ===== CONFIG AND CACHING
# =======================================================================

# Dùng st.secrets để quản lý khóa API một cách an toàn hơn
# st.secrets['GEMINI_API_KEY']
GEMINI_API_KEY = "AIzaSyCSenmJGRf2VJ9WId1SwQpfL3dMRRaHWmw"
# GEMINI_API_KEY = "AIzaSyABwa4KRue_M2A7l2YHAN4J2tPQJ5s33Ig"
# GEMINI_API_KEY = "AIzaSyBGSw2-NoZXd3HT_jWK1HoNzX7WhHcaBNA"
genai.configure(api_key=GEMINI_API_KEY)

# Sử dụng caching cho các tài nguyên nặng
@st.cache_resource
def load_embedding_model():
    return SentenceTransformer('sentence-transformers/all-MiniLM-L6-v2')

@st.cache_resource
def get_chroma_client():
    return chromadb.Client()

@st.cache_resource
def get_chroma_collection(_client):
    return _client.get_or_create_collection("essays")

embedding_model = load_embedding_model()
chroma_client = get_chroma_client()
collection = get_chroma_collection(chroma_client)

# =======================================================================
# ===== RUBRIC AND DATA LOADING
# =======================================================================

rubric = """
Tiêu chí chấm điểm thuộc lĩnh vực Công nghệ thông tin. Tổng điểm tối đa là 10 điểm, chia thành 5 nhóm tiêu chí chính:
1. Nội dung và kiến thức chuyên môn (4.5 điểm)
2. Cấu trúc và tổ chức bài viết (2 điểm)
3. Trình bày và diễn đạt (1.5 điểm)
4. Phân tích, lập luận và giải thích (1.5 điểm)
5. Tài liệu tham khảo và trích dẫn (0.5 điểm)
Ngoài ra có tiêu chí trừ điểm nếu bài làm vi phạm yêu cầu nghiêm trọng.

Tiêu chí 1 – Nội dung và kiến thức chuyên môn (4.5 điểm):
• Đáp ứng yêu cầu đề bài (1.5 điểm): bài làm đúng trọng tâm, không lạc đề, bao quát đủ các khía cạnh.
• Tính chính xác kiến thức (2 điểm): nội dung đúng, không sai nghiêm trọng về lý thuyết, thuật toán, phương pháp.
• Tính sáng tạo và thực tiễn (1 điểm): có tiếp cận mới, vận dụng thực tế, có dẫn chứng rõ.

Tiêu chí 2 – Cấu trúc và tổ chức bài viết (2 điểm):
• Cấu trúc (1.2 điểm): bài gồm mở bài, thân bài, kết luận; bố cục rõ ràng, không rời rạc.
• Logic liên kết (0.8 điểm): các ý nối liền mạch, có sử dụng từ nối để giữ tính liên tục.

Tiêu chí 3 – Trình bày và diễn đạt (1.5 điểm):
• Ngôn ngữ chuyên ngành (0.6 điểm): sử dụng đúng thuật ngữ CNTT, tránh cách nói đơn giản hóa.
• Chính tả, ngữ pháp (0.6 điểm): không mắc lỗi nghiêm trọng, diễn đạt mạch lạc.
• Hình thức trình bày (0.3 điểm): rõ ràng, khoa học, có số trang/mục lục nếu cần.

Tiêu chí 4 – Phân tích, lập luận và giải thích (1.5 điểm):
• Lập luận chặt chẽ (0.7 điểm): có cơ sở rõ ràng, giải thích hợp lý.
• Ví dụ, dẫn chứng (0.5 điểm): có minh họa bằng thực tế, sơ đồ hoặc thuật toán.
• Phản biện, đánh giá (0.3 điểm): đưa nhiều góc nhìn, so sánh giải pháp khác nhau.

Tiêu chí 5 – Tài liệu tham khảo và trích dẫn (0.5 điểm):
• Tài liệu chất lượng (0.3 điểm): có nguồn gốc rõ ràng, chính thống.
• Trích dẫn đúng chuẩn (0.2 điểm): theo quy chuẩn như APA, IEEE, Harvard...

Tiêu chí trừ điểm:
• Lạc đề hoặc sai nghiêm trọng về kiến thức: trừ toàn bộ điểm phần nội dung.
• Trình bày quá sơ sài, không có cấu trúc rõ ràng: trừ tối đa 2 điểm.
• Viết sai chính tả, ngữ pháp quá nhiều: trừ tối đa 1 điểm.
• Không trích nguồn nhưng dùng tài liệu ngoài: trừ 0.5 điểm.
"""

def export_collection_to_csv(collection, file_path="essays_backup.csv"):
    all_data = collection.get()
    with open(file_path, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow(["id", "document", "course", "category", "filename", "basename"])
        
        for i in range(len(all_data["ids"])):
            metadata = all_data["metadatas"][i] if "metadatas" in all_data else {}
            writer.writerow([
                all_data["ids"][i],
                all_data["documents"][i],
                metadata.get("course", ""),
                metadata.get("category", ""),
                metadata.get("filename", ""),
                metadata.get("basename", "")
            ])

# Sửa đổi hàm load_sample_data
def load_sample_data():
    base_folder = "data"
    docs, ids, metadatas = [], [], []

    if not os.path.exists(base_folder):
        os.makedirs(base_folder)

    for course in os.listdir(base_folder):
        course_path = os.path.join(base_folder, course)
        if not os.path.isdir(course_path):
            continue

        for category in ["essays", "scores", "teaching_materials"]:
            folder_path = os.path.join(course_path, category)
            if not os.path.exists(folder_path):
                os.makedirs(folder_path)

            for idx, filename in enumerate(os.listdir(folder_path)):
                path = os.path.join(folder_path, filename)
                text = ""
                
                if filename.endswith(".txt"):
                    with open(path, "r", encoding="utf-8") as f:
                        text = f.read()
                elif filename.endswith(".pdf"):
                    with fitz.open(path) as doc:
                        for page in doc:
                            page_text = page.get_text("text")
                            if page_text:
                                text += page_text + "\n"
                elif filename.endswith(".docx"):
                    doc = docx.Document(path)
                    text = "\n".join([para.text for para in doc.paragraphs])
                else:
                    continue

                if text.strip():
                    basename = os.path.splitext(filename)[0]
                    docs.append(text)
                    doc_id = f"{course}_{category}_{basename}"
                    ids.append(doc_id)
                    
                    metadata = {
                        "course": course,
                        "category": category,
                        "filename": filename,
                        "basename": basename
                    }

                    if category == "essays":
                        score_file_path = os.path.join(course_path, "scores", f"{basename}.score")
                        if os.path.exists(score_file_path):
                            try:
                                with open(score_file_path, "r", encoding="utf-8") as f:
                                    metadata["sample_score"] = float(f.read().strip())
                            except (IOError, ValueError):
                                print(f"Không thể đọc điểm từ file: {score_file_path}")

                    metadatas.append(metadata)
    
    if docs:
        embeddings = embedding_model.encode(docs).tolist()
        collection.add(documents=docs, ids=ids, embeddings=embeddings, metadatas=metadatas)
        export_collection_to_csv(collection, "essays_backup.csv")

if "data_loaded" not in st.session_state:
    load_sample_data()
    st.session_state["data_loaded"] = True


# =======================================================================
# ===== FUNCTIONS
# =======================================================================

def read_file(file):
    try:
        if file.name.endswith(".pdf"):
            import fitz
            text = ""
            with fitz.open(stream=file.read(), filetype="pdf") as doc:
                for page in doc:
                    page_text = page.get_text("text")
                    if page_text:
                        text += page_text + "\n"
            return text
        elif file.name.endswith(".docx"):
            doc = docx.Document(file)
            return "\n".join([para.text for para in doc.paragraphs])
        elif file.name.endswith(".txt"):
            return file.read().decode("utf-8")
        else:
            return ""
    except Exception as e:
        st.error(f"❌ Lỗi khi đọc file: {e}")
        return ""

def save_result_to_excel(course, filename, essay_text, score, file_path="grading_results.xlsx"):
    from openpyxl import Workbook, load_workbook
    
    essay_preview = essay_text[:300] + ("..." if len(essay_text) > 300 else "")
    
    if os.path.exists(file_path):
        wb = load_workbook(file_path)
        ws = wb.active
    else:
        wb = Workbook()
        ws = wb.active
        ws.append(["Tên môn học", "Tên file", "Nội dung bài luận", "Điểm"])

    ws.append([course, filename, essay_preview, score])
    wb.save(file_path)

def find_relevant_docs(query, course_context):
    query_embedding = embedding_model.encode([query]).tolist()
    
    # Thêm bộ lọc metadata vào truy vấn
    results = collection.query(
        query_embeddings=query_embedding,
        n_results=1,
        include=['metadatas', 'documents'],
        where={"course": course_context}
    )
    
    if results['documents'] and results['metadatas']:
        # Lấy tài liệu và metadata của tài liệu liên quan nhất
        relevant_doc = results['documents'][0][0]
        metadata = results['metadatas'][0][0]
        
        # Lấy điểm mẫu từ metadata (nếu có)
        sample_score = metadata.get("sample_score", None)
        
        return relevant_doc, sample_score
    return "", None

def grade_essay(essay_text, course_context, student_context, sample_text, sample_score):
    temperature_value = 0.5
    model = genai.GenerativeModel("gemini-2.0-flash")
    
    # Thêm thông tin điểm mẫu vào prompt
    sample_score_info = ""
    if sample_score is not None:
        sample_score_info = f"Bài luận mẫu này có điểm thực tế là: {sample_score} điểm. Hãy so sánh và giải thích lý do tại sao bài làm của sinh viên có thể cao hơn hoặc thấp hơn điểm này."
    else:
        sample_score_info = "Không có điểm mẫu tham chiếu."

    prompt = f"""
Bạn là một tiến sĩ đại học có nhiều năm kinh nghiệm trong lĩnh vực Công nghệ thông tin. Bạn cực kỳ nghiêm khắc và chuyên nghiệp trong việc chấm điểm, luôn đưa ra phản hồi chi tiết, khách quan và có tính xây dựng.

Hãy phân tích và chấm điểm bài luận của sinh viên dưới đây dựa trên bộ tiêu chí chi tiết mà tôi cung cấp. Bài luận này thuộc môn học {course_context}.

Bộ tiêu chí chấm điểm:
{rubric}

Yêu cầu cụ thể:
1.  Chấm điểm từng tiêu chí: Phân tích kỹ lưỡng và đưa ra điểm số cụ thể (có thể là số thập phân) cho từng tiêu chí trong bộ rubric. Sử dụng thang điểm từ 0 đến 10 một cách linh hoạt, không giới hạn điểm trong một khoảng hẹp.
2.  So sánh với bài mẫu: Đối chiếu bài làm của sinh viên với bài luận mẫu {sample_text} được cung cấp. Nhấn mạnh những điểm mà sinh viên đã làm tốt hơn hoặc chưa bằng bài mẫu.
3.  Nhận xét chi tiết:
        Điểm mạnh: Nêu rõ những mặt tích cực mà sinh viên đã làm được, ví dụ: "Kiến thức chính xác và lập luận chặt chẽ."
        Điểm cần cải thiện: Đề xuất những điểm mà sinh viên có thể cải thiện trong bài viết tiếp theo để đạt điểm cao hơn.
4.  Tính tổng điểm cuối cùng: Tổng hợp điểm từ các tiêu chí và đưa ra tổng điểm cuối cùng trên thang điểm 10.
5.  Thể hiện ngữ cảnh sinh viên: Dựa vào mô tả {student_context}, hãy điều chỉnh mức độ nghiêm khắc khi chấm điểm. Ví dụ, với sinh viên năm nhất, tập trung vào cấu trúc cơ bản; với sinh viên năm 4, yêu cầu cao hơn về mặt học thuật và phân tích.

Thông tin tham chiếu:
-   Điểm mẫu để tham khảo: {sample_score_info}
-   Bài luận mẫu cùng môn học: {sample_text}
-   Bài luận cần chấm: {essay_text}

Định dạng đầu ra:
Hãy trả về kết quả theo cấu trúc sau, đảm bảo mọi thông tin đều được trình bày rõ ràng và dễ đọc.
-   Điểm tổng: [Điểm số linh hoạt từ 0-10]
-   Điểm mạnh: [Đưa ra ít nhất 3 điểm mạnh nổi bật]
-   Điểm cần cải thiện: [Đưa ra ít nhất 3 điểm cần cải thiện cụ thể]
-   Nhận xét: [Phản hồi chi tiết theo yêu cầu trên]
"""
    try:
        response = model.generate_content(
            prompt,
            generation_config={"temperature": temperature_value}
        )
        return response.text
    except Exception as e:
        return f"Đã xảy ra lỗi khi chấm điểm: {e}"

def extract_score(result_text: str) -> str:
    score_value = "?"
    for line in result_text.splitlines():
        if any(kw in line for kw in ["Điểm tổng", "Tổng điểm", "Điểm cuối cùng", "Tổng điểm cuối cùng"]):
            match = re.search(r"(\d+(\.\d+)?)", line)
            if match:
                score_value = match.group(1)
            break

    if score_value == "?":
        matches = re.findall(r"(\d+(\.\d+)?)", result_text)
        candidates = [float(m[0]) for m in matches]
        candidates = [c for c in candidates if 0 <= c <= 10]
        if candidates:
            score_value = str(max(candidates))

    return score_value


# =======================================================================
# ===== STREAMLIT UI
# =======================================================================

st.set_page_config(page_title="Chấm điểm bài luận", page_icon="📄", layout="wide")

st.markdown("""
    <style>
    .main {
        background-color: #f9fafc;
    }
    .stTitle {
        font-size: 28px !important;
        color: #1a237e;
        font-weight: bold;
    }
    .score-box {
        padding: 20px;
        border-radius: 12px;
        background-color: #e3f2fd;
        text-align: center;
        margin-bottom: 20px;
    }
    .score-text {
        font-size: 48px;
        font-weight: bold;
        color: #0d47a1;
    }
    </style>
""", unsafe_allow_html=True)

st.title("📄 Hệ thống chấm điểm bài luận tự động")

col1, col2 = st.columns([1, 2])

with col1:
    st.header("📂 Nhập dữ liệu")
    uploaded_file = st.file_uploader("Tải lên bài luận (.pdf, .docx, .txt)", type=["pdf", "docx", "txt"])
    course_context = st.selectbox("📘 Môn học", ["Chiến lược kinh doanh thương mại điện tử", "Hệ thống thanh toán điện tử", "Marketing kỹ thuật số", "Phát triển ứng dụng thương mại điện tử", "Quản trị dự án thương mại điện tử", "Thương mại điện tử"])
    student_context = st.selectbox("🎓 Ngữ cảnh sinh viên", [
        "Năm 4: Sinh viên cần đạt chuẩn luận văn tốt nghiệp: viết học thuật hoàn chỉnh, có đặt vấn đề, cơ sở lý thuyết, phương pháp, phân tích, kết quả, kết luận, sử dụng trích dẫn chuẩn, thể hiện tư duy nghiên cứu độc lập và đóng góp mới.", 
        "Năm 3: Sinh viên phải thể hiện lập luận chặt chẽ hơn, biết sử dụng tài liệu tham khảo (trích dẫn đúng cách), trình bày theo chuẩn học thuật, có phân tích, so sánh, đánh giá thay vì chỉ mô tả.", 
        "Năm 2: Sinh viên bắt đầu học kỹ năng viết nâng cao hơn. Cần có cấu trúc 3 phần (mở bài, thân bài, kết luận), biết triển khai luận điểm theo đoạn văn mạch lạc, có ví dụ minh họa cơ bản.", 
        "Năm 1: Sinh viên chỉ mới làm quen viết luận. Đánh giá chủ yếu ở sự rõ ràng, logic cơ bản, cách trình bày ý tưởng. Không yêu cầu nhiều về trích dẫn học thuật hay cấu trúc phức tạp."
    ])

    # Trong phần xử lý nút bấm
    if st.button("🚀 Chấm điểm", use_container_width=True):
        if uploaded_file is not None:
            with st.status("🚀 Đang chấm điểm...", expanded=True) as status_box:
                status_box.write("Đang đọc nội dung file...")
                essay_text = read_file(uploaded_file)
                
                if not essay_text.strip():
                    status_box.update(label="❌ Lỗi: Không đọc được nội dung từ file.", state="error", expanded=False)
                    st.error("❌ Không đọc được nội dung từ file.")
                else:
                    status_box.write("Đang tìm tài liệu tham khảo...")
                    
                    # Gọi hàm find_relevant_docs và nhận cả 2 giá trị
                    sample_text, sample_score = find_relevant_docs(essay_text, course_context)

                    if not sample_text:
                        status_box.update(label="⚠️ Không tìm thấy bài mẫu nào cho môn học này.", state="warning", expanded=False)
                        st.warning(f"Không tìm thấy bài luận mẫu nào trong cơ sở dữ liệu cho môn học: {course_context}. Kết quả chấm điểm có thể không chính xác.")
                        
                        # Vẫn tiếp tục chấm điểm nhưng không có ngữ cảnh bài mẫu và điểm
                        result = grade_essay(
                            essay_text=essay_text, 
                            course_context=course_context, 
                            student_context=student_context, 
                            sample_text="Không có bài mẫu tham chiếu.", 
                            sample_score=None # Gửi None nếu không tìm thấy điểm
                        )
                        st.session_state["grading_result"] = result
                        score_value = extract_score(result)
                        
                        status_box.update(label=f"✅ Đã hoàn tất! Kết quả: {score_value} điểm.", state="complete", expanded=False)
                        save_result_to_excel(course_context, uploaded_file.name, essay_text, score_value)
                        st.success(f"✅ Kết quả đã được lưu vào grading_results.xlsx (Điểm AI: {score_value})")
                    else:
                        status_box.write("Đang phân tích bài luận với mô hình AI...")
                        
                        # Gọi hàm chấm điểm với tham số điểm mẫu
                        result = grade_essay(
                            essay_text=essay_text, 
                            course_context=course_context, 
                            student_context=student_context, 
                            sample_text=sample_text,
                            sample_score=sample_score
                        )
                        st.session_state["grading_result"] = result
                        
                        score_value = extract_score(result)

                        status_box.write(f"✅ Đã hoàn tất! Kết quả: {score_value} điểm.")
                        status_box.update(label="✅ Đã chấm điểm xong!", state="complete", expanded=False)
                        
                        save_result_to_excel(course_context, uploaded_file.name, essay_text, score_value)
                        st.success(f"✅ Kết quả đã được lưu vào grading_results.xlsx (Điểm AI: {score_value})")
        else:
            st.warning("⚠️ Vui lòng tải lên bài luận trước.")

with col2:
    st.header("📊 Kết quả chấm điểm")
    if "grading_result" in st.session_state:
        result_text = st.session_state["grading_result"]
        score_value = extract_score(result_text)
        st.markdown(f"""
            <div class="score-box">
                <div style="color: #000">Điểm tổng</div>
                <div class="score-text">{score_value}</div>
            </div>
        """, unsafe_allow_html=True)

        st.subheader("📝 Nhận xét chi tiết")
        st.write(result_text)
    else:
        st.info("👉 Kết quả sẽ hiển thị tại đây sau khi chấm điểm.")

# =======================================================================
# ===== QUẢN LÝ CORPUS
# =======================================================================
st.markdown("---")
st.header("📂 Quản lý Corpus")

if st.button("🔄 Làm mới dữ liệu corpus"):
    # Clear cache trước khi tải lại
    st.session_state.data_loaded = False
    st.cache_resource.clear()
    st.rerun()

all_data = collection.get()
if all_data and all_data["ids"]:
    import pandas as pd
    rows = []
    for i in range(len(all_data["ids"])):
        meta = all_data["metadatas"][i] if "metadatas" in all_data else {}
        rows.append({
            "id": all_data["ids"][i],
            "course": meta.get("course", ""),
            "category": meta.get("category", ""),
            "filename": meta.get("filename", ""),
            "basename": meta.get("basename", ""),
            "preview": all_data["documents"][i][:200] + "..."
        })
    df = pd.DataFrame(rows)
    st.dataframe(df, use_container_width=True)
else:
    st.info("📭 Chưa có dữ liệu corpus. Vui lòng thêm file vào thư mục `data/` và làm mới.")